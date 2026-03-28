from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_hms_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Hospital Management System – SQL Queries and Results', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    queries = [
        {
            "num": 1,
            "desc": "List all female patients over 20 years old.",
            "sql": "SELECT PatientID, Name, Age, Gender, Phone, Address\nFROM Patient\nWHERE Gender = 'Female' AND Age > 20;",
            "caption": "Figure 1: This screenshot shows the output of my query listing female patients older than 20 years."
        },
        {
            "num": 2,
            "desc": "List all available (Free) beds in the 'General' ward.",
            "sql": "SELECT BedNumber, WardType, Status\nFROM bed\nWHERE Status = 'Free' AND WardType = 'General';",
            "caption": "Figure 2: This screenshot shows the output of my query listing available beds in the General ward."
        },
        {
            "num": 3,
            "desc": "Display all staff members who are Doctors with their specialties and salaries.",
            "sql": "SELECT Name, Specialty, Salary\nFROM staff\nWHERE Role = 'Doctor';",
            "caption": "Figure 3: This screenshot shows the output of my query listing doctors and their professional details."
        },
        {
            "num": 4,
            "desc": "List all medicines with a price greater than 100, sorted by price.",
            "sql": "SELECT Medicine, Manufacturer, Price\nFROM medicine\nWHERE Price > 100\nORDER BY Price DESC;",
            "caption": "Figure 4: This screenshot shows the output of my query listing premium medicines."
        },
        {
            "num": 5,
            "desc": "Show patient names and their respective diagnoses from the medical records.",
            "sql": "SELECT p.Name AS PatientName, m.Diagnosis, m.AdmissionDate\nFROM Patient p\nINNER JOIN [Medical-record] m ON p.PatientID = m.PatientID;",
            "caption": "Figure 5: This screenshot shows the output of my query showing patient treatment history."
        },
        {
            "num": 6,
            "desc": "Search for all staff members who have 'Physician' in their specialty.",
            "sql": "SELECT Name, Specialty\nFROM staff\nWHERE Specialty LIKE '*Physician*';",
            "caption": "Figure 6: Using the * wildcard to find any specialty containing the word 'Physician'."
        },
        {
            "num": 7,
            "desc": "List patients whose phone numbers start with '0300'.",
            "sql": "SELECT Name, Phone\nFROM Patient\nWHERE Phone LIKE '0300*';",
            "caption": "Figure 7: Using the * wildcard at the end to match prefixes."
        },
        {
            "num": 8,
            "desc": "Find medicines that start with letters A through M.",
            "sql": "SELECT Medicine, Price\nFROM medicine\nWHERE Medicine LIKE '[A-M]*';",
            "caption": "Figure 8: Using character ranges [A-M] to filter results alphabetically."
        },
        {
            "num": 9,
            "desc": "Find all Doctors whose names end with 'an'.",
            "sql": "SELECT Name\nFROM staff\nWHERE Name LIKE 'Dr *an';",
            "caption": "Figure 9: Using the * wildcard at the beginning to match suffixes."
        },
        {
            "num": 10,
            "desc": "Search for all medicines that contain 'cin' in their name.",
            "sql": "SELECT Medicine, Price\nFROM medicine\nWHERE Medicine LIKE '*cin*';",
            "caption": "Figure 10: Using double wildcards to find a substring anywhere in the name."
        }
    ]

    for q in queries:
        # Query Label and Description
        p = doc.add_paragraph()
        run = p.add_run(f'Query {q["num"]}: {q["desc"]}')
        run.bold = True
        run.font.size = Pt(12)

        # SQL Box (simulated with shading or just a box)
        sql_p = doc.add_paragraph()
        sql_p.paragraph_format.left_indent = Inches(0.5)
        sql_run = sql_p.add_run(q["sql"])
        sql_run.font.name = 'Consolas'
        sql_run.font.size = Pt(10)

        # Space for screenshot
        doc.add_paragraph("\n[PASTE YOUR SCREENSHOT HERE]\n")

        # Caption
        caption = doc.add_paragraph()
        caption_run = caption.add_run(q["caption"])
        caption_run.italic = True
        caption_run.bold = True
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page separator
        doc.add_page_break()

    doc.save('HMSResults.docx')
    print("HMSResults.docx created successfully!")

if __name__ == "__main__":
    create_hms_doc()
